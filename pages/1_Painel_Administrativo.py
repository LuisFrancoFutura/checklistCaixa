import streamlit as st
import json
import os
import pandas as pd
import plotly.express as px
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.units import inch

# --- Configuração da Página ---
st.set_page_config(
    page_title="Painel Administrativo",
    page_icon="🛡️",
    layout="wide",
)

# --- Constante para o arquivo de histórico ---
COMPLETED_FILE = "completed_checklists.json"

# --- Funções de Persistência (duplicadas para modularidade) ---
def load_completed_tickets():
    if not os.path.exists(COMPLETED_FILE):
        return {}
    with open(COMPLETED_FILE, 'r', encoding='utf-8') as f:
        try:
            return json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            return {}

# --- Funções de Geração de Relatório (duplicadas para modularidade) ---
def get_report_data(ticket_data):
    agencia = ticket_data.get('agencia', '')
    cidade_uf = ticket_data.get('cidade_uf', '')
    endereco = ticket_data.get('endereco', '')
    num_racks = int(ticket_data.get('num_racks', 1))
    
    report_lines = ["TITLE: Check list Caixa Econômica", ""]
    report_lines.extend([f"Agência: {agencia}", f"Cidade/UF: {cidade_uf}", f"Endereço: {endereco}", f"Quantidade de Rack na agência: {num_racks}", ""])

    for i in range(1, num_racks + 1):
        report_lines.extend([f"SUBTITLE: Rack {i}:", f"Local instalado: {ticket_data.get(f'rack_local_{i}', '')}", f"Tamanho do Rack {i} – Número de Us: {ticket_data.get(f'rack_tamanho_{i}', '')}", f"Quantidade de Us disponíveis: {ticket_data.get(f'rack_us_disponiveis_{i}', '')}", f"Quantidade de réguas de energia: {ticket_data.get(f'rack_reguas_{i}', '')}", f"Quantidade de tomadas disponíveis: {ticket_data.get(f'rack_tomadas_disponiveis_{i}', '')}", f"Disponibilidade para ampliação de réguas de energia: {ticket_data.get(f'rack_ampliacao_reguas_{i}', 'Não')}", f"Rack está em bom estado: {ticket_data.get(f'rack_estado_{i}', 'Não')}", f"Rack está organizado: {ticket_data.get(f'rack_organizado_{i}', 'Não')}", f"Equipamentos e cabeamentos identificados: {ticket_data.get(f'rack_identificado_{i}', 'Não')}", ""])

    report_lines.extend(["SUBTITLE: Access Point (AP)", "", f"Verificar a quantidade de APs: {ticket_data.get('ap_quantidade', '')}", f"Identificar o setor onde será instalado*: {ticket_data.get('ap_setor', '')}", f"Verificar as condições da Instalação (se possui infra ou não): {ticket_data.get('ap_condicoes', '')}", f"** Altura que será instalado / distância do rack até o ponto de instalação: {ticket_data.get('ap_distancia', '')}"])
    
    return report_lines

def create_pdf_report(ticket_data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='Title', parent=styles['h1'], fontName='Helvetica-Bold', fontSize=14, alignment=TA_CENTER, spaceAfter=20)
    subtitle_style = ParagraphStyle(name='Subtitle', parent=styles['h2'], fontName='Helvetica-Bold', fontSize=12, alignment=TA_LEFT, spaceAfter=10)
    body_style = ParagraphStyle(name='Body', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, spaceAfter=4)
    story = []
    for line in get_report_data(ticket_data):
        if line.startswith("TITLE:"): story.append(Paragraph(line.replace("TITLE:", "").strip(), title_style))
        elif line.startswith("SUBTITLE:"): story.append(Paragraph(line.replace("SUBTITLE:", "").strip(), subtitle_style))
        elif line.strip() == "": story.append(Spacer(1, 0.1*inch))
        else: story.append(Paragraph(line.replace("<br>", "&nbsp;<br/>&nbsp;"), body_style))
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_report(ticket_data):
    document = Document()
    for line in get_report_data(ticket_data):
        if line.startswith("TITLE:"):
            p = document.add_paragraph(); p.add_run(line.replace("TITLE:", "").strip()).bold = True; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("SUBTITLE:"):
            p = document.add_paragraph(); p.add_run(line.replace("SUBTITLE:", "").strip()).bold = True
        else: document.add_paragraph(line)
    buffer = io.BytesIO(); document.save(buffer); buffer.seek(0)
    return buffer

# --- Funções de Exibição da UI ---
def display_review_checklist(ticket_id, data_source):
    """Renderiza o formulário em modo de leitura."""
    
    with st.expander("Informações Gerais da Agência", expanded=True):
        st.text(f"Agência: {data_source.get('agencia', 'N/A')}")
        st.text(f"Cidade/UF: {data_source.get('cidade_uf', 'N/A')}")
        st.text(f"Endereço: {data_source.get('endereco', 'N/A')}")
        st.text(f"Quantidade de Racks: {int(data_source.get('num_racks', 1))}")

    num_racks = int(data_source.get('num_racks', 1))

    with st.expander("Detalhes dos Racks"):
        for i in range(1, num_racks + 1):
            st.markdown(f"#### Rack {i}")
            st.text(f"Local: {data_source.get(f'rack_local_{i}', 'N/A')}")
            st.text(f"Tamanho (U's): {data_source.get(f'rack_tamanho_{i}', 'N/A')}")
            st.text(f"U's disponíveis: {data_source.get(f'rack_us_disponiveis_{i}', 'N/A')}")
            st.text(f"Réguas de energia: {data_source.get(f'rack_reguas_{i}', 'N/A')}")
            st.text(f"Tomadas disponíveis: {data_source.get(f'rack_tomadas_disponiveis_{i}', 'N/A')}")
            st.text(f"Permite ampliação de réguas: {data_source.get(f'rack_ampliacao_reguas_{i}', 'N/A')}")
            st.text(f"Bom estado: {data_source.get(f'rack_estado_{i}', 'N/A')}")
            st.text(f"Organizado: {data_source.get(f'rack_organizado_{i}', 'N/A')}")
            st.text(f"Identificado: {data_source.get(f'rack_identificado_{i}', 'N/A')}")
            if i < num_racks: st.markdown("---")
            
    with st.expander("Access Point (AP)"):
        st.text(f"APs existentes: {data_source.get('ap_quantidade', 'N/A')}")
        st.text(f"Setor de instalação: {data_source.get('ap_setor', 'N/A')}")
        st.text(f"Condições da infra: {data_source.get('ap_condicoes', 'N/A')}")
        st.text(f"Altura/Distância: {data_source.get('ap_distancia', 'N/A')}")

    st.markdown("---")
    st.subheader("Exportar Relatório")
    d_col1, d_col2, d_col3 = st.columns(3)
    with d_col1: st.download_button("Baixar .TXT", "\n".join(get_report_data(data_source)), f"Checklist_{ticket_id.upper()}.txt", "text/plain")
    with d_col2: st.download_button("Baixar .PDF", create_pdf_report(data_source), f"Checklist_{ticket_id.upper()}.pdf", "application/pdf")
    with d_col3: st.download_button("Baixar .DOCX", create_docx_report(data_source), f"Checklist_{ticket_id.upper()}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# --- Telas do Admin ---
def page_admin_login():
    st.header("Login do Painel Administrativo")
    with st.form("login_form"):
        username = st.text_input("Usuário")
        password = st.text_input("Senha", type="password")
        submitted = st.form_submit_button("Login")
        if submitted:
            if username == "admin" and password == "admin":
                st.session_state.logged_in = True
                st.rerun()
            else:
                st.error("Usuário ou senha incorretos")

def page_admin_dashboard():
    st.title("Painel Administrativo")
    tab1, tab2 = st.tabs(["Revisão de Chamados", "Estatísticas"])

    with tab1:
        st.header("Revisar Chamados Concluídos")
        completed_tickets = load_completed_tickets()
        if not completed_tickets:
            st.info("Nenhum chamado concluído para revisar.")
        else:
            options = ["Selecione..."] + list(completed_tickets.keys())
            ticket_to_review = st.selectbox("Selecione um chamado:", options=options, key="review_select")
            if ticket_to_review != "Selecione...":
                st.subheader(f"Revisando Chamado: {ticket_to_review.upper()}")
                display_review_checklist(ticket_to_review, completed_tickets[ticket_to_review])

    with tab2:
        st.header("Estatísticas dos Checklists")
        completed_tickets = load_completed_tickets()
        if not completed_tickets:
            st.warning("Não há dados de chamados concluídos para gerar estatísticas.")
        else:
            df = pd.DataFrame.from_dict(completed_tickets, orient='index')
            st.metric("Total de Chamados Concluídos", len(df))

            st.subheader("Chamados por Localização (Cidade/UF)")
            location_counts = df['cidade_uf'].value_counts().reset_index()
            location_counts.columns = ['Localização', 'Contagem']
            fig_loc = px.bar(location_counts, x='Localização', y='Contagem', title="Distribuição de Chamados")
            st.plotly_chart(fig_loc, use_container_width=True)

            st.subheader("Análise de Status dos Racks")
            status_keys = {'estado': 'Rack em bom estado', 'organizado': 'Rack organizado', 'identificado': 'Equipamentos identificados'}
            status_counts = {key: {'Sim': 0, 'Não': 0} for key in status_keys}
            
            for _, ticket_data in df.iterrows():
                num_racks = int(ticket_data.get('num_racks', 1))
                for i in range(1, num_racks + 1):
                    for key, _ in status_keys.items():
                        status_val = ticket_data.get(f'rack_{key}_{i}', 'Não')
                        if status_val in ['Sim', 'Não']: status_counts[key][status_val] += 1

            col1, col2, col3 = st.columns(3)
            with col1:
                fig1 = px.pie(values=list(status_counts['estado'].values()), names=list(status_counts['estado'].keys()), title=status_keys['estado'])
                st.plotly_chart(fig1, use_container_width=True)
            with col2:
                fig2 = px.pie(values=list(status_counts['organizado'].values()), names=list(status_counts['organizado'].keys()), title=status_keys['organizado'])
                st.plotly_chart(fig2, use_container_width=True)
            with col3:
                fig3 = px.pie(values=list(status_counts['identificado'].values()), names=list(status_counts['identificado'].keys()), title=status_keys['identificado'])
                st.plotly_chart(fig3, use_container_width=True)


# --- Lógica Principal de Navegação ---
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

if st.session_state.logged_in:
    page_admin_dashboard()
else:
    page_admin_login()
