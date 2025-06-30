import streamlit as st
import json
import os
import pandas as pd
import plotly.express as px
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.units import inch

# --- Constante para o arquivo de histórico ---
COMPLETED_FILE = "completed_checklists.json"

# --- CSS Melhorado (mesmo do app principal) ---
def load_admin_css():
    """Carrega e injeta o CSS customizado melhorado para o painel administrativo."""
    css = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

        /* Reset e configurações base */
        * {
            box-sizing: border-box;
        }

        html {
            font-size: 16px;
        }

        body, .main {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            background-color: #f8fafc !important;
            line-height: 1.6 !important;
            color: #1e293b !important;
        }

        /* Oculta a barra lateral do Streamlit */
        [data-testid="stSidebar"] {
            display: none !important;
        }

        /* Container principal */
        .main .block-container {
            padding-top: 2rem !important;
            padding-bottom: 2rem !important;
            padding-left: 1rem !important;
            padding-right: 1rem !important;
            max-width: 1200px !important;
        }

        /* Títulos melhorados */
        h1, h2, h3, h4, h5, h6 {
            color: #0f172a !important;
            font-weight: 600 !important;
            margin-bottom: 1rem !important;
            line-height: 1.2 !important;
        }

        h1 {
            font-size: 2.25rem !important;
            margin-bottom: 1.5rem !important;
        }

        h2 {
            font-size: 1.875rem !important;
            margin-bottom: 1.25rem !important;
        }

        h3 {
            font-size: 1.5rem !important;
        }

        h4 {
            font-size: 1.25rem !important;
        }

        /* Botões melhorados */
        .stButton > button, .stDownloadButton > button {
            background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
            color: white !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 12px 24px !important;
            font-weight: 500 !important;
            font-size: 1rem !important;
            min-height: 48px !important;
            width: 100% !important;
            transition: all 0.2s ease !important;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06) !important;
        }

        .stButton > button:hover, .stDownloadButton > button:hover {
            background: linear-gradient(135deg, #2563eb 0%, #1e40af 100%) !important;
            transform: translateY(-1px) !important;
            box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1), 0 4px 6px -2px rgba(0, 0, 0, 0.05) !important;
        }

        /* Botão de logout com cor diferente */
        .stButton > button[kind="secondary"] {
            background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%) !important;
        }

        .stButton > button[kind="secondary"]:hover {
            background: linear-gradient(135deg, #dc2626 0%, #b91c1c 100%) !important;
        }

        /* Expanders melhorados */
        [data-testid="stExpander"] {
            background-color: #ffffff !important;
            border: 1px solid #e2e8f0 !important;
            border-radius: 16px !important;
            margin-bottom: 1.5rem !important;
            box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1), 0 1px 2px 0 rgba(0, 0, 0, 0.06) !important;
            overflow: hidden !important;
        }

        [data-testid="stExpander"] summary {
            font-size: 1.125rem !important;
            font-weight: 600 !important;
            color: #0f172a !important;
            padding: 1.25rem 1.5rem !important;
            background-color: #f8fafc !important;
            border-bottom: 1px solid #e2e8f0 !important;
            cursor: pointer !important;
        }

        [data-testid="stExpander"] summary:hover {
            background-color: #f1f5f9 !important;
        }

        [data-testid="stExpander"] > div:last-child {
            padding: 1.5rem !important;
        }

        /* Labels melhorados */
        [data-testid="stWidgetLabel"] label {
            color: #374151 !important;
            font-weight: 500 !important;
            font-size: 0.95rem !important;
            margin-bottom: 0.5rem !important;
        }

        /* Inputs melhorados */
        [data-testid="stTextInput"] input, 
        [data-testid="stNumberInput"] input,
        [data-testid="stTextArea"] textarea,
        [data-testid="stSelectbox"] select {
            border: 2px solid #e2e8f0 !important;
            border-radius: 8px !important;
            padding: 12px 16px !important;
            font-size: 1rem !important;
            min-height: 48px !important;
            background-color: #ffffff !important;
            transition: border-color 0.2s ease !important;
        }

        [data-testid="stTextInput"] input:focus, 
        [data-testid="stNumberInput"] input:focus,
        [data-testid="stTextArea"] textarea:focus,
        [data-testid="stSelectbox"] select:focus {
            border-color: #3b82f6 !important;
            box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1) !important;
            outline: none !important;
        }

        /* Formulários melhorados */
        [data-testid="stForm"] {
            background-color: #ffffff !important;
            padding: 2rem !important;
            border-radius: 16px !important;
            border: 1px solid #e2e8f0 !important;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06) !important;
        }

        /* Tabs melhorados */
        [data-testid="stTabs"] {
            margin-bottom: 2rem !important;
        }

        [data-testid="stTabs"] button {
            font-size: 1rem !important;
            font-weight: 500 !important;
            padding: 12px 24px !important;
            border-radius: 8px 8px 0 0 !important;
        }

        /* Métricas melhoradas */
        [data-testid="metric-container"] {
            background-color: #ffffff !important;
            padding: 1.5rem !important;
            border-radius: 12px !important;
            border: 1px solid #e2e8f0 !important;
            box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1) !important;
        }

        /* Alertas e mensagens */
        [data-testid="stAlert"] {
            border-radius: 12px !important;
            padding: 1rem 1.5rem !important;
            margin: 1rem 0 !important;
        }

        /* Texto de informação melhorado */
        .info-text {
            background-color: #ffffff !important;
            padding: 1rem !important;
            border-radius: 8px !important;
            border-left: 4px solid #3b82f6 !important;
            margin: 0.5rem 0 !important;
        }

        /* Responsividade para mobile */
        @media (max-width: 768px) {
            html {
                font-size: 14px;
            }

            .main .block-container {
                padding-top: 1rem !important;
                padding-left: 0.75rem !important;
                padding-right: 0.75rem !important;
            }

            h1 {
                font-size: 1.75rem !important;
                text-align: center !important;
            }

            h2 {
                font-size: 1.5rem !important;
            }

            /* Colunas empilhadas em mobile */
            [data-testid="stHorizontalBlock"] {
                flex-direction: column !important;
                gap: 1rem !important;
            }

            [data-testid="stHorizontalBlock"] > div {
                width: 100% !important;
                margin-bottom: 0 !important;
            }

            /* Botões maiores em mobile */
            .stButton > button, .stDownloadButton > button {
                min-height: 52px !important;
                font-size: 1.05rem !important;
                padding: 16px 24px !important;
            }

            /* Inputs maiores em mobile */
            [data-testid="stTextInput"] input, 
            [data-testid="stNumberInput"] input,
            [data-testid="stTextArea"] textarea,
            [data-testid="stSelectbox"] select {
                min-height: 52px !important;
                font-size: 1.05rem !important;
                padding: 16px !important;
            }

            /* Gráficos responsivos */
            [data-testid="stPlotlyChart"] {
                width: 100% !important;
                overflow-x: auto !important;
            }

            /* Formulários com menos padding em mobile */
            [data-testid="stForm"] {
                padding: 1.5rem 1rem !important;
            }
        }

        /* Responsividade para telas muito pequenas */
        @media (max-width: 480px) {
            html {
                font-size: 13px;
            }

            .main .block-container {
                padding-left: 0.5rem !important;
                padding-right: 0.5rem !important;
            }

            h1 {
                font-size: 1.5rem !important;
            }
        }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

# --- Funções de Persistência (duplicadas para modularidade) ---
def load_completed_tickets():
    if not os.path.exists(COMPLETED_FILE):
        return {}
    with open(COMPLETED_FILE, 'r', encoding='utf-8') as f:
        try:
            return json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            return {}

# --- Funções de Geração de Relatório (duplicadas para modularidade) ---
def get_report_data(ticket_data):
    agencia = ticket_data.get('agencia', '')
    cidade_uf = ticket_data.get('cidade_uf', '')
    endereco = ticket_data.get('endereco', '')
    num_racks = int(ticket_data.get('num_racks', 1))
    
    report_lines = ["TITLE: Check list Caixa Econômica", ""]
    report_lines.extend([f"Agência: {agencia}", f"Cidade/UF: {cidade_uf}", f"Endereço: {endereco}", f"Quantidade de Rack na agência: {num_racks}", ""])

    for i in range(1, num_racks + 1):
        report_lines.extend([f"SUBTITLE: Rack {i}:", f"Local instalado: {ticket_data.get(f'rack_local_{i}', '')}", f"Tamanho do Rack {i} – Número de Us: {ticket_data.get(f'rack_tamanho_{i}', '')}", f"Quantidade de Us disponíveis: {ticket_data.get(f'rack_us_disponiveis_{i}', '')}", f"Quantidade de réguas de energia: {ticket_data.get(f'rack_reguas_{i}', '')}", f"Quantidade de tomadas disponíveis: {ticket_data.get(f'rack_tomadas_disponiveis_{i}', '')}", f"Disponibilidade para ampliação de réguas de energia: {ticket_data.get(f'rack_ampliacao_reguas_{i}', 'Não')}", f"Rack está em bom estado: {ticket_data.get(f'rack_estado_{i}', 'Não')}", f"Rack está organizado: {ticket_data.get(f'rack_organizado_{i}', 'Não')}", f"Equipamentos e cabeamentos identificados: {ticket_data.get(f'rack_identificado_{i}', 'Não')}", ""])

    report_lines.extend(["SUBTITLE: Access Point (AP)", "", f"Verificar a quantidade de APs: {ticket_data.get('ap_quantidade', '')}", f"Identificar o setor onde será instalado*: {ticket_data.get('ap_setor', '')}", f"Verificar as condições da Instalação (se possui infra ou não): {ticket_data.get('ap_condicoes', '')}", f"** Altura que será instalado / distância do rack até o ponto de instalação: {ticket_data.get('ap_distancia', '')}"])
    
    return report_lines

def create_pdf_report(ticket_data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='Title', parent=styles['h1'], fontName='Helvetica-Bold', fontSize=14, alignment=TA_CENTER, spaceAfter=20)
    subtitle_style = ParagraphStyle(name='Subtitle', parent=styles['h2'], fontName='Helvetica-Bold', fontSize=12, alignment=TA_LEFT, spaceAfter=10)
    body_style = ParagraphStyle(name='Body', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, spaceAfter=4)
    story = []
    for line in get_report_data(ticket_data):
        if line.startswith("TITLE:"): story.append(Paragraph(line.replace("TITLE:", "").strip(), title_style))
        elif line.startswith("SUBTITLE:"): story.append(Paragraph(line.replace("SUBTITLE:", "").strip(), subtitle_style))
        elif line.strip() == "": story.append(Spacer(1, 0.1*inch))
        else: story.append(Paragraph(line.replace("<br>", "&nbsp;<br/>&nbsp;"), body_style))
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_report(ticket_data):
    document = Document()
    for line in get_report_data(ticket_data):
        if line.startswith("TITLE:"):
            p = document.add_paragraph(); p.add_run(line.replace("TITLE:", "").strip()).bold = True; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("SUBTITLE:"):
            p = document.add_paragraph(); p.add_run(line.replace("SUBTITLE:", "").strip()).bold = True
        else: document.add_paragraph(line)
    buffer = io.BytesIO(); document.save(buffer); buffer.seek(0)
    return buffer

# --- Funções de Exibição da UI do Admin ---
def display_review_checklist(ticket_id, data_source):
    """Renderiza o formulário em modo de leitura."""
    
    with st.expander("📋 Informações Gerais da Agência", expanded=True):
        col1, col2 = st.columns([1, 1])
        with col1:
            st.markdown(f"**🏢 Agência:** {data_source.get('agencia', 'N/A')}")
            st.markdown(f"**📍 Endereço:** {data_source.get('endereco', 'N/A')}")
        with col2:
            st.markdown(f"**🌍 Cidade/UF:** {data_source.get('cidade_uf', 'N/A')}")
            st.markdown(f"**🗄️ Quantidade de Racks:** {int(data_source.get('num_racks', 1))}")

    num_racks = int(data_source.get('num_racks', 1))

    with st.expander("🗄️ Detalhes dos Racks", expanded=True):
        for i in range(1, num_racks + 1):
            st.markdown(f"#### 📦 Rack {i}")
            col1, col2 = st.columns([1, 1])
            with col1:
                st.markdown(f"**📍 Local:** {data_source.get(f'rack_local_{i}', 'N/A')}")
                st.markdown(f"**📏 Tamanho (U's):** {data_source.get(f'rack_tamanho_{i}', 'N/A')}")
                st.markdown(f"**📊 U's disponíveis:** {data_source.get(f'rack_us_disponiveis_{i}', 'N/A')}")
                st.markdown(f"**⚡ Réguas de energia:** {data_source.get(f'rack_reguas_{i}', 'N/A')}")
                st.markdown(f"**🔌 Tomadas disponíveis:** {data_source.get(f'rack_tomadas_disponiveis_{i}', 'N/A')}")
            with col2:
                st.markdown(f"**🔧 Permite ampliação de réguas:** {data_source.get(f'rack_ampliacao_reguas_{i}', 'N/A')}")
                st.markdown(f"**✅ Bom estado:** {data_source.get(f'rack_estado_{i}', 'N/A')}")
                st.markdown(f"**🗂️ Organizado:** {data_source.get(f'rack_organizado_{i}', 'N/A')}")
                st.markdown(f"**🏷️ Identificado:** {data_source.get(f'rack_identificado_{i}', 'N/A')}")
            if i < num_racks: st.markdown("---")
            
    with st.expander("📡 Access Point (AP)", expanded=True):
        col1, col2 = st.columns([1, 1])
        with col1:
            st.markdown(f"**📊 APs existentes:** {data_source.get('ap_quantidade', 'N/A')}")
            st.markdown(f"**🎯 Setor de instalação:** {data_source.get('ap_setor', 'N/A')}")
        with col2:
            st.markdown(f"**🔍 Condições da infra:** {data_source.get('ap_condicoes', 'N/A')}")
            st.markdown(f"**📐 Altura/Distância:** {data_source.get('ap_distancia', 'N/A')}")

    st.markdown("---")
    st.subheader("📄 Exportar Relatório")
    d_col1, d_col2, d_col3 = st.columns(3)
    with d_col1: 
        st.download_button("📄 Baixar .TXT", "\n".join(get_report_data(data_source)), f"Checklist_{ticket_id.upper()}.txt", "text/plain")
    with d_col2: 
        st.download_button("📑 Baixar .PDF", create_pdf_report(data_source), f"Checklist_{ticket_id.upper()}.pdf", "application/pdf")
    with d_col3: 
        st.download_button("📝 Baixar .DOCX", create_docx_report(data_source), f"Checklist_{ticket_id.upper()}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# --- Telas do Admin ---
def page_admin_login():
    load_admin_css()
    
    st.title("🔐 Painel Administrativo")
    st.header("🔑 Login")
    
    with st.form("login_form"):
        st.markdown("### 👤 Credenciais de Acesso")
        username = st.text_input("👤 Usuário", placeholder="Digite seu usuário")
        password = st.text_input("🔒 Senha", type="password", placeholder="Digite sua senha")
        
        col1, col2 = st.columns([1, 1])
        with col1:
            submitted = st.form_submit_button("🚀 Entrar", type="primary")
        with col2:
            back_button = st.form_submit_button("⬅️ Voltar")
            
        if submitted:
            if username == "admin" and password == "admin":
                st.session_state.logged_in = True
                st.session_state.page = 'admin_dashboard'
                st.success("✅ Login realizado com sucesso!")
                st.rerun()
            else:
                st.error("❌ Usuário ou senha incorretos")
                
        if back_button:
            st.session_state.page = 'main'
            st.rerun()

def page_admin_dashboard():
    load_admin_css()
    
    if not st.session_state.get('logged_in'):
        st.session_state.page = 'admin_login'
        st.error("❌ Acesso negado. Por favor, faça o login.")
        st.rerun()
    
    st.title("📊 Painel Administrativo")
    
    # Header com botão de logout
    col_title, col_logout = st.columns([3, 1])
    with col_logout:
        if st.button("🚪 Sair", type="secondary"):
            st.session_state.page = 'main'
            if 'logged_in' in st.session_state:
                del st.session_state.logged_in
            st.rerun()

    tab1, tab2 = st.tabs(["📋 Revisão de Chamados", "📈 Estatísticas"])

    with tab1:
        st.header("🔍 Revisar Chamados Concluídos")
        completed_tickets = load_completed_tickets()
        
        if not completed_tickets:
            st.info("ℹ️ Nenhum chamado concluído para revisar.")
        else:
            st.success(f"✅ {len(completed_tickets)} chamados encontrados")
            
            options = ["Selecione um chamado..."] + list(completed_tickets.keys())
            ticket_to_review = st.selectbox(
                "🎫 Selecione um chamado:", 
                options=options, 
                key="review_select"
            )
            
            if ticket_to_review != "Selecione um chamado...":
                st.subheader(f"📋 Revisando Chamado: {ticket_to_review.upper()}")
                display_review_checklist(ticket_to_review, completed_tickets[ticket_to_review])

    with tab2:
        st.header("📊 Estatísticas dos Checklists")
        completed_tickets = load_completed_tickets()
        
        if not completed_tickets:
            st.warning("⚠️ Não há dados de chamados concluídos para gerar estatísticas.")
        else:
            df = pd.DataFrame.from_dict(completed_tickets, orient='index')
            
            # Métricas principais
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📊 Total de Chamados", len(df))
            with col2:
                total_racks = df['num_racks'].astype(int).sum()
                st.metric("🗄️ Total de Racks", total_racks)
            with col3:
                avg_racks = df['num_racks'].astype(int).mean()
                st.metric("📈 Média de Racks/Chamado", f"{avg_racks:.1f}")

            st.markdown("---")

            # Gráfico de distribuição por localização
            st.subheader("🌍 Chamados por Localização (Cidade/UF)")
            if 'cidade_uf' in df.columns and not df['cidade_uf'].isna().all():
                location_counts = df['cidade_uf'].value_counts().reset_index()
                location_counts.columns = ['Localização', 'Contagem']
                
                fig_loc = px.bar(
                    location_counts, 
                    x='Localização', 
                    y='Contagem', 
                    title="📍 Distribuição de Chamados por Localização",
                    color='Contagem',
                    color_continuous_scale='Blues'
                )
                fig_loc.update_layout(
                    xaxis_title="Localização",
                    yaxis_title="Número de Chamados",
                    showlegend=False
                )
                st.plotly_chart(fig_loc, use_container_width=True)
            else:
                st.info("ℹ️ Dados de localização não disponíveis")

            st.markdown("---")

            # Análise de status dos racks
            st.subheader("🔍 Análise de Status dos Racks")
            status_keys = {
                'estado': '✅ Rack em bom estado', 
                'organizado': '🗂️ Rack organizado', 
                'identificado': '🏷️ Equipamentos identificados'
            }
            status_counts = {key: {'Sim': 0, 'Não': 0} for key in status_keys}
            
            for _, ticket_data in df.iterrows():
                num_racks = int(ticket_data.get('num_racks', 1))
                for i in range(1, num_racks + 1):
                    for key, _ in status_keys.items():
                        status_val = ticket_data.get(f'rack_{key}_{i}', 'Não')
                        if status_val in ['Sim', 'Não']: 
                            status_counts[key][status_val] += 1

            col1, col2, col3 = st.columns(3)
            
            with col1:
                if sum(status_counts['estado'].values()) > 0:
                    fig1 = px.pie(
                        values=list(status_counts['estado'].values()), 
                        names=list(status_counts['estado'].keys()), 
                        title=status_keys['estado'],
                        color_discrete_sequence=['#10b981', '#ef4444']
                    )
                    st.plotly_chart(fig1, use_container_width=True)
                else:
                    st.info("Sem dados")
                    
            with col2:
                if sum(status_counts['organizado'].values()) > 0:
                    fig2 = px.pie(
                        values=list(status_counts['organizado'].values()), 
                        names=list(status_counts['organizado'].keys()), 
                        title=status_keys['organizado'],
                        color_discrete_sequence=['#10b981', '#ef4444']
                    )
                    st.plotly_chart(fig2, use_container_width=True)
                else:
                    st.info("Sem dados")
                    
            with col3:
                if sum(status_counts['identificado'].values()) > 0:
                    fig3 = px.pie(
                        values=list(status_counts['identificado'].values()), 
                        names=list(status_counts['identificado'].keys()), 
                        title=status_keys['identificado'],
                        color_discrete_sequence=['#10b981', '#ef4444']
                    )
                    st.plotly_chart(fig3, use_container_width=True)
                else:
                    st.info("Sem dados")

