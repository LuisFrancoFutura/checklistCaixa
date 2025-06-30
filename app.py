# --- INSTRUÇÕES DE INSTALAÇÃO ---
# Antes de rodar, instale as bibliotecas necessárias:
# pip install streamlit python-docx reportlab pandas plotly

import streamlit as st
import datetime
import io
import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.units import inch

# --- Configuração da Página ---
st.set_page_config(
    page_title="Checklist Help Desk",
    page_icon="✨",
    layout="wide",
)

# --- Constante para o arquivo de histórico ---
COMPLETED_FILE = "completed_checklists.json"

# --- CSS Melhorado para um Design Responsivo e Legível ---
def load_css():
    """Carrega e injeta o CSS customizado melhorado para estilizar a aplicação."""
    css = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

        /* Reset e configurações base */
        * {
            box-sizing: border-box;
        }

        html {
            font-size: 16px;
        }

        body, .main {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
            background-color: #f8fafc !important;
            line-height: 1.6 !important;
            color: #1e293b !important;
        }

        /* Oculta a barra lateral do Streamlit */
        [data-testid="stSidebar"] {
            display: none !important;
        }

        /* Container principal */
        .main .block-container {
            padding-top: 2rem !important;
            padding-bottom: 2rem !important;
            padding-left: 1rem !important;
            padding-right: 1rem !important;
            max-width: 1200px !important;
        }

        /* Títulos melhorados */
        h1, h2, h3, h4, h5, h6 {
            color: #0f172a !important;
            font-weight: 600 !important;
            margin-bottom: 1rem !important;
            line-height: 1.2 !important;
        }

        h1 {
            font-size: 2.25rem !important;
            margin-bottom: 1.5rem !important;
        }

        h2 {
            font-size: 1.875rem !important;
            margin-bottom: 1.25rem !important;
        }

        h3 {
            font-size: 1.5rem !important;
        }

        h4 {
            font-size: 1.25rem !important;
        }

        /* Botões melhorados */
        .stButton > button, .stDownloadButton > button {
            background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
            color: white !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 12px 24px !important;
            font-weight: 500 !important;
            font-size: 1rem !important;
            min-height: 48px !important;
            width: 100% !important;
            transition: all 0.2s ease !important;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06) !important;
        }

        .stButton > button:hover, .stDownloadButton > button:hover {
            background: linear-gradient(135deg, #2563eb 0%, #1e40af 100%) !important;
            transform: translateY(-1px) !important;
            box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1), 0 4px 6px -2px rgba(0, 0, 0, 0.05) !important;
        }

        .stButton > button:active, .stDownloadButton > button:active {
            transform: translateY(0) !important;
        }

        /* Botão primário especial */
        .stButton > button[kind="primary"] {
            background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
        }

        .stButton > button[kind="primary"]:hover {
            background: linear-gradient(135deg, #059669 0%, #047857 100%) !important;
        }

        /* Expanders melhorados */
        [data-testid="stExpander"] {
            background-color: #ffffff !important;
            border: 1px solid #e2e8f0 !important;
            border-radius: 16px !important;
            margin-bottom: 1.5rem !important;
            box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1), 0 1px 2px 0 rgba(0, 0, 0, 0.06) !important;
            overflow: hidden !important;
        }

        [data-testid="stExpander"] summary {
            font-size: 1.125rem !important;
            font-weight: 600 !important;
            color: #0f172a !important;
            padding: 1.25rem 1.5rem !important;
            background-color: #f8fafc !important;
            border-bottom: 1px solid #e2e8f0 !important;
            cursor: pointer !important;
        }

        [data-testid="stExpander"] summary:hover {
            background-color: #f1f5f9 !important;
        }

        [data-testid="stExpander"] > div:last-child {
            padding: 1.5rem !important;
        }

        /* Labels melhorados */
        [data-testid="stWidgetLabel"] label {
            color: #374151 !important;
            font-weight: 500 !important;
            font-size: 0.95rem !important;
            margin-bottom: 0.5rem !important;
        }

        /* Inputs melhorados */
        [data-testid="stTextInput"] input, 
        [data-testid="stNumberInput"] input,
        [data-testid="stTextArea"] textarea {
            border: 2px solid #e2e8f0 !important;
            border-radius: 8px !important;
            padding: 12px 16px !important;
            font-size: 1rem !important;
            min-height: 48px !important;
            background-color: #ffffff !important;
            transition: border-color 0.2s ease !important;
        }

        [data-testid="stTextInput"] input:focus, 
        [data-testid="stNumberInput"] input:focus,
        [data-testid="stTextArea"] textarea:focus {
            border-color: #3b82f6 !important;
            box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1) !important;
            outline: none !important;
        }

        /* Radio buttons melhorados */
        [data-testid="stRadio"] {
            margin-bottom: 1rem !important;
        }

        [data-testid="stRadio"] > div {
            gap: 1rem !important;
        }

        [data-testid="stRadio"] label {
            font-size: 0.95rem !important;
            padding: 8px 16px !important;
            border: 2px solid #e2e8f0 !important;
            border-radius: 8px !important;
            background-color: #ffffff !important;
            cursor: pointer !important;
            transition: all 0.2s ease !important;
            min-height: 44px !important;
            display: flex !important;
            align-items: center !important;
        }

        [data-testid="stRadio"] label:hover {
            border-color: #3b82f6 !important;
            background-color: #f8fafc !important;
        }

        /* Formulários melhorados */
        [data-testid="stForm"] {
            background-color: #ffffff !important;
            padding: 2rem !important;
            border-radius: 16px !important;
            border: 1px solid #e2e8f0 !important;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06) !important;
        }

        /* Tabs melhorados */
        [data-testid="stTabs"] {
            margin-bottom: 2rem !important;
        }

        [data-testid="stTabs"] button {
            font-size: 1rem !important;
            font-weight: 500 !important;
            padding: 12px 24px !important;
            border-radius: 8px 8px 0 0 !important;
        }

        /* Métricas melhoradas */
        [data-testid="metric-container"] {
            background-color: #ffffff !important;
            padding: 1.5rem !important;
            border-radius: 12px !important;
            border: 1px solid #e2e8f0 !important;
            box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1) !important;
        }

        /* Separadores */
        hr {
            border: none !important;
            height: 1px !important;
            background-color: #e2e8f0 !important;
            margin: 2rem 0 !important;
        }

        /* Alertas e mensagens */
        [data-testid="stAlert"] {
            border-radius: 12px !important;
            padding: 1rem 1.5rem !important;
            margin: 1rem 0 !important;
        }

        /* Responsividade para tablets */
        @media (max-width: 1024px) {
            .main .block-container {
                padding-left: 1rem !important;
                padding-right: 1rem !important;
            }

            h1 {
                font-size: 2rem !important;
            }

            h2 {
                font-size: 1.75rem !important;
            }
        }

        /* Responsividade para mobile */
        @media (max-width: 768px) {
            html {
                font-size: 14px;
            }

            .main .block-container {
                padding-top: 1rem !important;
                padding-left: 0.75rem !important;
                padding-right: 0.75rem !important;
            }

            h1 {
                font-size: 1.75rem !important;
                text-align: center !important;
            }

            h2 {
                font-size: 1.5rem !important;
            }

            h3 {
                font-size: 1.25rem !important;
            }

            /* Colunas empilhadas em mobile */
            [data-testid="stHorizontalBlock"] {
                flex-direction: column !important;
                gap: 1rem !important;
            }

            [data-testid="stHorizontalBlock"] > div {
                width: 100% !important;
                margin-bottom: 0 !important;
            }

            /* Botões maiores em mobile */
            .stButton > button, .stDownloadButton > button {
                min-height: 52px !important;
                font-size: 1.05rem !important;
                padding: 16px 24px !important;
            }

            /* Inputs maiores em mobile */
            [data-testid="stTextInput"] input, 
            [data-testid="stNumberInput"] input,
            [data-testid="stTextArea"] textarea {
                min-height: 52px !important;
                font-size: 1.05rem !important;
                padding: 16px !important;
            }

            /* Expanders com menos padding em mobile */
            [data-testid="stExpander"] summary {
                padding: 1rem !important;
                font-size: 1rem !important;
            }

            [data-testid="stExpander"] > div:last-child {
                padding: 1rem !important;
            }

            /* Radio buttons em coluna em mobile */
            [data-testid="stRadio"] > div {
                flex-direction: column !important;
                gap: 0.75rem !important;
            }

            [data-testid="stRadio"] label {
                width: 100% !important;
                justify-content: center !important;
                min-height: 48px !important;
            }

            /* Formulários com menos padding em mobile */
            [data-testid="stForm"] {
                padding: 1.5rem 1rem !important;
            }

            /* Gráficos responsivos */
            [data-testid="stPlotlyChart"] {
                width: 100% !important;
                overflow-x: auto !important;
            }
        }

        /* Responsividade para telas muito pequenas */
        @media (max-width: 480px) {
            html {
                font-size: 13px;
            }

            .main .block-container {
                padding-left: 0.5rem !important;
                padding-right: 0.5rem !important;
            }

            h1 {
                font-size: 1.5rem !important;
            }

            [data-testid="stExpander"] summary {
                padding: 0.75rem !important;
            }

            [data-testid="stExpander"] > div:last-child {
                padding: 0.75rem !important;
            }
        }

        /* Estados de foco melhorados para acessibilidade */
        button:focus-visible,
        input:focus-visible,
        textarea:focus-visible {
            outline: 2px solid #3b82f6 !important;
            outline-offset: 2px !important;
        }

        /* Animações suaves */
        * {
            transition: background-color 0.2s ease, border-color 0.2s ease, color 0.2s ease !important;
        }

        /* Melhor contraste para texto */
        p, span, div {
            color: #374151 !important;
        }

        /* Espaçamento consistente */
        .element-container {
            margin-bottom: 1rem !important;
        }

        /* Loading states */
        [data-testid="stSpinner"] {
            color: #3b82f6 !important;
        }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

# --- Funções de Persistência ---
def load_completed_tickets():
    if not os.path.exists(COMPLETED_FILE):
        return {}
    with open(COMPLETED_FILE, 'r', encoding='utf-8') as f:
        try:
            return json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            return {}

def save_completed_ticket(ticket_id, data):
    all_completed = load_completed_tickets()
    all_completed[ticket_id] = data
    with open(COMPLETED_FILE, 'w', encoding='utf-8') as f:
        json.dump(all_completed, f, indent=4, ensure_ascii=False)


# --- Funções de Geração de Relatório ---
def get_report_data(ticket_data):
    agencia = ticket_data.get('agencia', '')
    cidade_uf = ticket_data.get('cidade_uf', '')
    endereco = ticket_data.get('endereco', '')
    num_racks = int(ticket_data.get('num_racks', 1))
    
    report_lines = ["TITLE: Check list Caixa Econômica", ""]
    report_lines.extend([f"Agência: {agencia}", f"Cidade/UF: {cidade_uf}", f"Endereço: {endereco}", f"Quantidade de Rack na agência: {num_racks}", ""])

    for i in range(1, num_racks + 1):
        report_lines.extend([f"SUBTITLE: Rack {i}:", f"Local instalado: {ticket_data.get(f'rack_local_{i}', '')}", f"Tamanho do Rack {i} – Número de Us: {ticket_data.get(f'rack_tamanho_{i}', '')}", f"Quantidade de Us disponíveis: {ticket_data.get(f'rack_us_disponiveis_{i}', '')}", f"Quantidade de réguas de energia: {ticket_data.get(f'rack_reguas_{i}', '')}", f"Quantidade de tomadas disponíveis: {ticket_data.get(f'rack_tomadas_disponiveis_{i}', '')}", f"Disponibilidade para ampliação de réguas de energia: {ticket_data.get(f'rack_ampliacao_reguas_{i}', 'Não')}", f"Rack está em bom estado: {ticket_data.get(f'rack_estado_{i}', 'Não')}", f"Rack está organizado: {ticket_data.get(f'rack_organizado_{i}', 'Não')}", f"Equipamentos e cabeamentos identificados: {ticket_data.get(f'rack_identificado_{i}', 'Não')}", ""])

    report_lines.extend(["SUBTITLE: Access Point (AP)", "", f"Verificar a quantidade de APs: {ticket_data.get('ap_quantidade', '')}", f"Identificar o setor onde será instalado*: {ticket_data.get('ap_setor', '')}", f"Verificar as condições da Instalação (se possui infra ou não): {ticket_data.get('ap_condicoes', '')}", f"** Altura que será instalado / distância do rack até o ponto de instalação: {ticket_data.get('ap_distancia', '')}"])
    
    return report_lines

def create_pdf_report(ticket_data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='Title', parent=styles['h1'], fontName='Helvetica-Bold', fontSize=14, alignment=TA_CENTER, spaceAfter=20)
    subtitle_style = ParagraphStyle(name='Subtitle', parent=styles['h2'], fontName='Helvetica-Bold', fontSize=12, alignment=TA_LEFT, spaceAfter=10)
    body_style = ParagraphStyle(name='Body', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, spaceAfter=4)
    story = []
    for line in get_report_data(ticket_data):
        if line.startswith("TITLE:"): story.append(Paragraph(line.replace("TITLE:", "").strip(), title_style))
        elif line.startswith("SUBTITLE:"): story.append(Paragraph(line.replace("SUBTITLE:", "").strip(), subtitle_style))
        elif line.strip() == "": story.append(Spacer(1, 0.1*inch))
        else: story.append(Paragraph(line.replace("<br>", "&nbsp;<br/>&nbsp;"), body_style))
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_report(ticket_data):
    document = Document()
    for line in get_report_data(ticket_data):
        if line.startswith("TITLE:"):
            p = document.add_paragraph(); p.add_run(line.replace("TITLE:", "").strip()).bold = True; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("SUBTITLE:"):
            p = document.add_paragraph(); p.add_run(line.replace("SUBTITLE:", "").strip()).bold = True
        else: document.add_paragraph(line)
    buffer = io.BytesIO(); document.save(buffer); buffer.seek(0)
    return buffer

# --- Funções de Exibição da UI ---
def display_checklist_form(ticket_id):
    """Renderiza os campos do formulário para um determinado chamado."""
    
    with st.expander("📋 Informações Gerais da Agência", expanded=True):
        col1, col2 = st.columns([1, 1])
        with col1:
            st.text_input("🏢 Agência", key=f'agencia_{ticket_id}', placeholder="Digite o nome da agência")
            st.text_input("📍 Endereço", key=f'endereco_{ticket_id}', placeholder="Endereço completo")
        with col2:
            st.text_input("🌍 Cidade/UF", key=f'cidade_uf_{ticket_id}', placeholder="Ex: São Paulo/SP")
            st.number_input("🗄️ Quantidade de Racks na agência", min_value=1, step=1, key=f'num_racks_{ticket_id}')
    
    num_racks = int(st.session_state.get(f'num_racks_{ticket_id}', 1))

    with st.expander("🗄️ Detalhes dos Racks", expanded=True):
        for i in range(1, num_racks + 1):
            st.markdown(f"#### 📦 Rack {i}")
            c1, c2 = st.columns([1, 1])
            with c1:
                st.text_input(f"📍 Local instalado", key=f'rack_local_{i}_{ticket_id}', placeholder="Ex: Sala de TI")
                st.text_input(f"📏 Tamanho do Rack {i} – Número de Us", key=f'rack_tamanho_{i}_{ticket_id}', placeholder="Ex: 42U")
                st.text_input(f"📊 Quantidade de Us disponíveis", key=f'rack_us_disponiveis_{i}_{ticket_id}', placeholder="Ex: 15U")
                st.text_input(f"⚡ Quantidade de réguas de energia", key=f'rack_reguas_{i}_{ticket_id}', placeholder="Ex: 2")
                st.text_input(f"🔌 Quantidade de tomadas disponíveis", key=f'rack_tomadas_disponiveis_{i}_{ticket_id}', placeholder="Ex: 8")
            with c2:
                radio_options = ("Sim", "Não")
                st.radio("🔧 Disponibilidade para ampliação de réguas de energia", radio_options, key=f'rack_ampliacao_reguas_{i}_{ticket_id}', horizontal=True)
                st.radio("✅ Rack está em bom estado", radio_options, key=f'rack_estado_{i}_{ticket_id}', horizontal=True)
                st.radio("🗂️ Rack está organizado", radio_options, key=f'rack_organizado_{i}_{ticket_id}', horizontal=True)
                st.radio("🏷️ Equipamentos e cabeamentos identificados", radio_options, key=f'rack_identificado_{i}_{ticket_id}', horizontal=True)
            if i < num_racks: st.markdown("---")

    with st.expander("📡 Access Point (AP)", expanded=True):
        st.text_input("📊 Verificar a quantidade de APs", key=f'ap_quantidade_{ticket_id}', placeholder="Ex: 5")
        st.text_input("🎯 Identificar o setor onde será instalado*", key=f'ap_setor_{ticket_id}', placeholder="Ex: Recepção, Gerência")
        st.text_input("🔍 Verificar as condições da Instalação", key=f'ap_condicoes_{ticket_id}', placeholder="Possui infra ou não")
        st.text_input("📐 ** Altura que será instalado / distância do rack", key=f'ap_distancia_{ticket_id}', placeholder="Ex: 3m altura / 15m distância")
    
    st.markdown("---")
    st.subheader("🎯 Ações")
    
    col_action1, col_action2 = st.columns([2, 1])
    with col_action1:
        if st.button("✅ Concluir e Arquivar Chamado", key=f"complete_{ticket_id}", type="primary"):
            current_ticket_data = {key.replace(f"_{ticket_id}", ""): value for key, value in st.session_state.items() if str(key).endswith(f"_{ticket_id}")}
            save_completed_ticket(ticket_id, current_ticket_data)
            st.session_state.active_ticket_id = None
            st.success(f"✅ Chamado {ticket_id} arquivado com sucesso!")
            st.rerun()
    
    with col_action2:
        if st.button("🔄 Iniciar outro chamado", key=f"new_ticket_{ticket_id}"):
            st.session_state.active_ticket_id = None
            st.rerun()

    final_ticket_data = {key.replace(f"_{ticket_id}", ""): value for key, value in st.session_state.items() if str(key).endswith(f"_{ticket_id}")}
    
    st.markdown("### 📄 Exportar Relatório")
    d_col1, d_col2, d_col3 = st.columns(3)
    with d_col1: 
        st.download_button("📄 Baixar .TXT", "\n".join(get_report_data(final_ticket_data)), f"Checklist_{ticket_id.upper()}.txt", "text/plain")
    with d_col2: 
        st.download_button("📑 Baixar .PDF", create_pdf_report(final_ticket_data), f"Checklist_{ticket_id.upper()}.pdf", "application/pdf")
    with d_col3: 
        st.download_button("📝 Baixar .DOCX", create_docx_report(final_ticket_data), f"Checklist_{ticket_id.upper()}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# --- Lógica Principal da Aplicação ---
load_css()

# Inicializa o estado da sessão
if 'active_ticket_id' not in st.session_state:
    st.session_state.active_ticket_id = None

st.title("🛠️ Ferramenta de Checklist de Campo")

if st.session_state.active_ticket_id is None:
    st.header("🚀 Iniciar Novo Checklist")
    
    # Botão para acessar painel administrativo
    col_main, col_admin = st.columns([3, 1])
    with col_admin:
        if st.button("🔐 Login ADM"):
            st.session_state.page = 'admin_login'
            st.rerun()
    
    with st.form("new_ticket_form"):
        st.markdown("### 🎫 Informações do Chamado")
        ticket_id_input = st.text_input("🔢 Insira o código do chamado:", placeholder="Ex: 12345 ou CLAR-12345")
        submitted = st.form_submit_button("🚀 Iniciar Checklist", type="primary")
        
        if submitted and ticket_id_input:
            formatted_id = ticket_id_input.strip().upper()
            if formatted_id.isdigit(): 
                formatted_id = f"CLAR-{formatted_id}"
            elif not formatted_id.startswith("CLAR-"): 
                formatted_id = f"CLAR-{formatted_id}"
            
            st.session_state.active_ticket_id = formatted_id
            for key in list(st.session_state.keys()):
                if key.endswith(f"_{formatted_id}"): 
                    del st.session_state[key]
            st.rerun()
else:
    ticket_id = st.session_state.active_ticket_id
    st.header(f"📋 Preenchendo Chamado: {ticket_id}")
    display_checklist_form(ticket_id)

# --- Lógica de Navegação para Admin ---
if st.session_state.get('page') == 'admin_login':
    from admin_page import page_admin_login
    page_admin_login()
elif st.session_state.get('page') == 'admin_dashboard':
    from admin_page import page_admin_dashboard
    page_admin_dashboard()

